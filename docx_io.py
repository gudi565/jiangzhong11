"""Word (.docx) I/O — 纯文本导出 + 专业版可编辑报告。

- extract_paragraphs / build_docx: 旧的「改写后纯文本」导出（make-docx 用）。
- build_report(task, orig_text, result): 5 类任务的可编辑 Word 报告（make-report 用），
  排版已算好的结果，逐句标色 + 原文/改写后双栏对照 + 指标摘要 + 出处表 + 免责声明。
"""
import io
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

CJK_FONT = "宋体"
LATIN_FONT = "Times New Roman"

COLOR_RGB = {
    "err": RGBColor(0xC0, 0x00, 0x00),    # 红
    "warn": RGBColor(0xB8, 0x6A, 0x00),   # 橙
    "ok": RGBColor(0x1A, 0x7A, 0x32),     # 绿
    "dark": RGBColor(0x22, 0x22, 0x22),   # 正文
    "mute": RGBColor(0x70, 0x70, 0x70),   # 辅助/免责
    "accent": RGBColor(0x0E, 0x6B, 0x4A), # 品牌深绿
}


def extract_paragraphs(file_bytes: bytes) -> list:
    doc = Document(io.BytesIO(file_bytes))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]


def _set_font(run, size=12):
    run.font.name = LATIN_FONT
    run.font.size = Pt(size)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), CJK_FONT)


def build_docx(paragraphs: list) -> bytes:
    doc = Document()
    for text in paragraphs:
        run = doc.add_paragraph().add_run(text)
        _set_font(run)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── 报告样式 helper ──────────────────────────────────────────────────────
def _add_run(p, text, color="dark", bold=False, size=12):
    r = p.add_run(text)
    r.font.color.rgb = COLOR_RGB.get(color, COLOR_RGB["dark"])
    r.bold = bold
    _set_font(r, size)
    return r


def _heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    sizes = {0: 18, 1: 14, 2: 12}
    _add_run(p, text, color=("accent" if level == 0 else "dark"),
             bold=True, size=sizes.get(level, 12))
    return p


def _metric_line(doc, pairs):
    """指标摘要行：pairs = [(label, value), ...] → 「标签: 值　标签: 值」"""
    p = doc.add_paragraph()
    for i, (label, value) in enumerate(pairs):
        if i:
            _add_run(p, "　", color="mute")
        _add_run(p, f"{label}: ", color="mute")
        _add_run(p, str(value), color="dark", bold=True)
    return p


def _hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p


def _disclaimer(doc, text):
    p = doc.add_paragraph()
    _add_run(p, text, color="mute", size=10)
    return p


def _colored_sentence(p, item):
    """item = {sentence, color, score/overlap?}。句尾标分 + 着色。"""
    color = item.get("color", "dark")
    s = item.get("sentence", "")
    tag = item.get("score", item.get("overlap"))
    _add_run(p, s, color=color)
    if tag is not None:
        _add_run(p, f" [{tag}]", color=color, size=10)
    _add_run(p, "。")


def _set_cell_width(cell, inches):
    cell.width = Inches(inches)


def _two_col_table(doc, left_title, right_title, orig_sents, right_items, right_hint=None):
    """原文|改写后 双栏对照表。right_items = [{sentence, color, ...}]。
    左列逐句排，右列按 right_items 逐句着色，按索引对齐，不足留空。"""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    _set_cell_width(hdr[0], 3.0)
    _set_cell_width(hdr[1], 3.0)
    hp0 = hdr[0].paragraphs[0]
    _add_run(hp0, left_title, bold=True, color="accent")
    hp1 = hdr[1].paragraphs[0]
    label = right_title + (f"（{right_hint}）" if right_hint else "")
    _add_run(hp1, label, bold=True, color="accent")

    n = max(len(orig_sents), len(right_items))
    for i in range(n):
        cells = table.add_row().cells
        _set_cell_width(cells[0], 3.0)
        _set_cell_width(cells[1], 3.0)
        lp = cells[0].paragraphs[0]
        left = orig_sents[i] if i < len(orig_sents) else ""
        _add_run(lp, left, color="dark")
        rp = cells[1].paragraphs[0]
        item = right_items[i] if i < len(right_items) else None
        if item:
            _colored_sentence(rp, item)
    return table


def _signals_table(doc, signals):
    """指标解读表：指标 | 数值 | 解读。signals = [{name, value, score, hint}]"""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for cell, title, w in [(hdr[0], "指标", 1.6), (hdr[1], "数值", 1.4), (hdr[2], "解读", 3.0)]:
        _set_cell_width(cell, w)
        _add_run(cell.paragraphs[0], title, bold=True, color="accent")
    for s in signals:
        cells = table.add_row().cells
        _set_cell_width(cells[0], 1.6)
        _set_cell_width(cells[1], 1.4)
        _set_cell_width(cells[2], 3.0)
        _add_run(cells[0].paragraphs[0], s.get("name", ""), bold=True)
        _add_run(cells[1].paragraphs[0], str(s.get("value", "")))
        _add_run(cells[2].paragraphs[0], s.get("hint", ""))
    return table


def _sources_table(doc, matches):
    """命中来源表：句子 | 雷同% | 出处。matches = [{sentence, overlap, title, url}]"""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for cell, title, w in [(hdr[0], "命中句子", 3.0), (hdr[1], "雷同%", 0.8), (hdr[2], "出处", 2.2)]:
        _set_cell_width(cell, w)
        _add_run(cell.paragraphs[0], title, bold=True, color="accent")
    for m in matches:
        cells = table.add_row().cells
        _set_cell_width(cells[0], 3.0)
        _set_cell_width(cells[1], 0.8)
        _set_cell_width(cells[2], 2.2)
        ov = m.get("overlap", 0)
        col = "err" if ov >= 50 else ("warn" if ov >= 30 else "ok")
        _add_run(cells[0].paragraphs[0], m.get("sentence", ""), color=col)
        _add_run(cells[1].paragraphs[0], f"{ov}%", color=col, bold=True)
        title = m.get("title", "") or "—"
        url = m.get("url", "")
        cp = cells[2].paragraphs[0]
        _add_run(cp, title, color="mute", size=10)
        if url:
            cp2 = cells[2].add_paragraph()
            _add_run(cp2, url, color="mute", size=9)
    return table


def _split_sents(text):
    import re
    return [s.strip() for s in re.split(r"[。！？\n;；]+", text or "") if s.strip()]


def _brand_header(doc, title, subtitle):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _add_run(p, "降重工具", color="accent", bold=True, size=12)
    _add_run(p, "  ·  ", color="mute")
    _add_run(p, title, bold=True, size=18)
    p2 = doc.add_paragraph()
    _add_run(p2, f"{subtitle}　生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')}",
             color="mute", size=10)
    _hr(doc)


# ── 总入口 ──────────────────────────────────────────────────────────────
def build_report(task: str, orig_text: str, result: dict) -> bytes:
    doc = Document()
    try:
        _build_dispatch[task](doc, orig_text or "", result or {})
    except KeyError:
        _heading(doc, "不支持的任务类型", 1)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _report_rewrite(doc, orig_text, result, title, right_hint):
    """降重 / 降AIGC / 英文修改 共用：双栏对照 + 指标 + 诊断。"""
    _brand_header(doc, title, f"强度 {result.get('strength', '—')} · 模式 {result.get('mode', '—')}")
    rw = result.get("rewrite", "")
    sim = result.get("similarity")
    cov = result.get("coverage")
    pairs = [
        ("改写后字数", result.get("len", len(rw))),
        ("相似度", f"{round(sim*100)}%" if isinstance(sim, (int, float)) else "—"),
        ("差异度", f"{round(cov*100)}%" if isinstance(cov, (int, float)) else "—"),
    ]
    if result.get("discipline"):
        pairs.append(("学科", result["discipline"]))
    if result.get("chunks") is not None:
        pairs.append(("分段", result["chunks"]))
    _metric_line(doc, pairs)
    _hr(doc)

    _heading(doc, "原文 / 改写后 对照（右列可在 Word 里直接继续修改）", 2)
    sentence_scores = result.get("sentence_scores") or []
    # 若没有逐句分数，把整段当一个 ok 项
    if not sentence_scores and rw:
        sentence_scores = [{"sentence": rw, "color": "ok"}]
    right_items = sentence_scores if sentence_scores else [{"sentence": rw, "color": "dark"}]
    orig_sents = _split_sents(orig_text) if orig_text else [orig_text]
    if not orig_sents:
        orig_sents = ["（无原文）"]
    left_title = "中文原文" if result.get("sub") == "translate" else "原文"
    right_title = "英文" if result.get("sub") == "translate" else "改写后"
    _two_col_table(doc, left_title, right_title, orig_sents, right_items, right_hint=right_hint)

    # 诊断附录（仅 pipeline 降重管线有）
    diag = result.get("diagnostics")
    stages = result.get("stages")
    if diag or stages:
        _hr(doc)
        _heading(doc, "改写诊断", 2)
        if stages:
            sp = doc.add_paragraph()
            _add_run(sp, "阶段：", bold=True, color="mute")
            _add_run(sp, " → ".join(stages))
        if isinstance(diag, list):
            for d in diag:
                p = doc.add_paragraph()
                _add_run(p, "· ", color="mute")
                lbl = d.get("label", "") if isinstance(d, dict) else str(d)
                _add_run(p, lbl)
                if isinstance(d, dict) and d.get("repaired"):
                    _add_run(p, "（已自动修复）", color="ok")
                if isinstance(d, dict) and d.get("still_missing"):
                    _add_run(p, f"　仍缺：{', '.join(d['still_missing'])}", color="warn", size=10)

    _hr(doc)
    _disclaimer(doc, "免责声明：本报告由 AI 改写引擎生成，差异度/相似度为估算值，仅供参考。"
                     "改写后内容仍需作者自行核对学术准确性、术语一致性及引用规范，"
                     "使用前请以学校或期刊的官方查重结果为准。")


def _report_aigc(doc, orig_text, result):
    score = result.get("aigc_score", 0)
    verdict = result.get("verdict", "")
    _brand_header(doc, "AIGC 检测报告", f"AIGC 疑似度 {score} · {verdict}")
    _metric_line(doc, [
        ("AIGC 疑似度", f"{score}/100"),
        ("判定", verdict),
        ("句数", result.get("sentence_count", "—")),
        ("字数", result.get("char_count", "—")),
        ("困惑度", result.get("perplexity", "—")),
        ("突发性", result.get("burstiness", "—")),
    ])
    _hr(doc)

    signals = result.get("signals") or []
    if signals:
        _heading(doc, "检测信号解读", 2)
        _signals_table(doc, signals)
        _hr(doc)

    _heading(doc, "正文逐句 AIGC 疑似度（红=偏AI / 橙=不确定 / 绿=偏人，可在 Word 里继续修改）", 2)
    sentence_scores = result.get("sentence_scores") or []
    if not sentence_scores:
        sentence_scores = [{"sentence": orig_text, "color": "warn", "score": score}]
    for item in sentence_scores:
        p = doc.add_paragraph()
        _colored_sentence(p, item)

    _hr(doc)
    _disclaimer(doc, result.get("note", "") +
                " AIGC 检测器普遍存在误判，本结果仅为参考指引，不作为定论。")


def _report_plagiarism(doc, orig_text, result):
    score = result.get("similarity_score", 0)
    verdict = result.get("verdict", "")
    _brand_header(doc, "查重报告", f"相似度估算 {score} · {verdict}")
    _metric_line(doc, [
        ("相似度估算", f"{score}/100"),
        ("判定", verdict),
        ("命中句数", result.get("matched_count", "—")),
        ("抽检句数", result.get("checked_count", "—")),
    ])
    _hr(doc)

    _heading(doc, "正文逐句相似度（红=高度雷同，重点改这些；可在 Word 里继续修改）", 2)
    sentence_scores = result.get("sentence_scores") or []
    if not sentence_scores:
        sentence_scores = [{"sentence": orig_text, "color": "warn", "overlap": score}]
    for item in sentence_scores:
        p = doc.add_paragraph()
        _colored_sentence(p, item)

    matches = result.get("matches") or []
    if matches:
        _hr(doc)
        _heading(doc, "命中来源明细", 2)
        _sources_table(doc, matches)

    _hr(doc)
    _disclaimer(doc, result.get("note", "") +
                " 本报告相似度为综合估算，接近但不等于知网精确查重率，仅供参考。"
                "权威查重率请以 cx.cnki.net 官方结果为准。")


_build_dispatch = {
    "rewrite": lambda d, o, r: _report_rewrite(d, o, r, "降重报告", "红=改得不够"),
    "humanize": lambda d, o, r: _report_rewrite(d, o, r, "降AIGC报告", "红=AI痕迹仍重"),
    "english": lambda d, o, r: _report_rewrite(d, o, r, "英文修改报告", "红=改动过小"),
    "aigc": _report_aigc,
    "plagiarism": _report_plagiarism,
}
